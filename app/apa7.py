from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


INCHES_1 = Inches(1)
CM_254 = Cm(2.54)


def configure_page(doc):
    for section in doc.sections:
        section.top_margin = CM_254
        section.bottom_margin = CM_254
        section.left_margin = CM_254
        section.right_margin = CM_254


def set_paragraph_format(paragraph, font_name='Times New Roman', font_size=12,
                         bold=False, italic=False, alignment=None,
                         space_before=0, space_after=0,
                         line_spacing=2.0, first_line_indent=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if alignment is not None:
        pf.alignment = alignment
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        _set_run_font_east_asia(run, font_name)


def _set_run_font_east_asia(run, font_name):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def configure_normal_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 2.0
    _set_run_font_east_asia_style(style, 'Times New Roman')
    _add_first_line_indent_to_style(style)


def _set_run_font_east_asia_style(style, font_name):
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _add_first_line_indent_to_style(style):
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style.element.append(pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), str(int(0.5 * 1440)))


def apply_apa7_heading_style(doc, level):
    style_name = f'Heading {level}'
    if style_name not in [s.name for s in doc.styles]:
        return
    style = doc.styles[style_name]
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = None
    _set_run_font_east_asia_style(style, 'Times New Roman')
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 2.0
    pf.first_line_indent = Inches(0)
    if level == 1:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.font.bold = True
        style.font.italic = False
    elif level == 2:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.font.bold = True
        style.font.italic = False
    elif level == 3:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.font.bold = True
        style.font.italic = True
    elif level == 4:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Inches(0.5)
        style.font.bold = True
        style.font.italic = False
    elif level == 5:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Inches(0.5)
        style.font.bold = True
        style.font.italic = True


def add_apa7_cover(doc, metadata):
    metadata = metadata or {}
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             first_line_indent=Inches(0))

    title = doc.add_paragraph()
    run = title.add_run(metadata.get('title', 'Title'))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    set_paragraph_format(title, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         first_line_indent=Inches(0))

    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             first_line_indent=Inches(0))

    fields = [
        metadata.get('author', 'Author Name'),
        metadata.get('institution', 'Institution'),
        metadata.get('course', 'Course'),
        metadata.get('professor', 'Professor'),
        metadata.get('date', 'Date'),
    ]
    for field_text in fields:
        p = doc.add_paragraph()
        run = p.add_run(field_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             first_line_indent=Inches(0))

    doc.add_page_break()


def add_page_numbers(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run()
        _add_page_number_field(run)


def _add_page_number_field(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), '24')


def add_reference_entry(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 2.0
    pf.first_line_indent = Inches(0)
    _set_hanging_indent(p, Inches(0.5))


def _set_hanging_indent(paragraph, indent_size):
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), str(int(indent_size.emu)))
    ind.set(qn('w:hanging'), str(int(indent_size.emu)))
