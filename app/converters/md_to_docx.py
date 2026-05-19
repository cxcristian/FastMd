import os
import re
import yaml
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.apa7 import (
    configure_page, configure_normal_style, apply_apa7_heading_style,
    add_apa7_cover, add_page_numbers, add_reference_entry,
    set_paragraph_format
)


class MdToDocxConverter:
    def __init__(self, md_path):
        self.md_path = md_path
        self.doc = None
        self.metadata = {}
        self.in_references = False
        self.image_dir = None
        self.image_counter = 0

    def convert(self, output_path, include_cover=True, include_page_numbers=True):
        with open(self.md_path, 'r', encoding='utf-8') as f:
            text = f.read()

        self.doc = Document()
        self.image_dir = os.path.join(os.path.dirname(output_path), 'images')
        self.doc = Document()

        configure_page(self.doc)
        configure_normal_style(self.doc)

        for i in range(1, 6):
            apply_apa7_heading_style(self.doc, i)

        blocks = self._parse_blocks(text)
        metadata = self._extract_front_matter(blocks)
        self.metadata = metadata
        content_blocks = [b for b in blocks if b[0] != 'front_matter']

        if include_cover:
            add_apa7_cover(self.doc, metadata)
        if include_page_numbers:
            add_page_numbers(self.doc)

        for block in content_blocks:
            self._render_block(block)

        self.doc.save(output_path)
        return True, 'Conversión APA 7 exitosa', self.metadata

    def _parse_blocks(self, text):
        blocks = []
        lines = text.split('\n')
        i = 0
        in_fence = False
        fence_lang = ''
        fence_content = []

        while i < len(lines):
            line = lines[i]

            if not in_fence and line.strip().startswith('```'):
                in_fence = True
                fence_lang = line.strip()[3:].strip()
                fence_content = []
                i += 1
                continue

            if in_fence:
                if line.strip().startswith('```'):
                    in_fence = False
                    blocks.append(('code', (fence_lang, '\n'.join(fence_content))))
                    i += 1
                    continue
                fence_content.append(line)
                i += 1
                continue

            if i == 0 and line.strip() == '---':
                end = self._find_line_index(lines, i + 1, lambda l: l.strip() == '---')
                if end is not None:
                    blocks.append(('front_matter', '\n'.join(lines[i + 1:end])))
                    i = end + 1
                    continue

            heading = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading:
                level = len(heading.group(1))
                text_content = heading.group(2)
                blocks.append(('heading', (level, text_content)))
                i += 1
                continue

            if not line.strip():
                i += 1
                continue

            hr_match = re.match(r'^[-*_]{3,}\s*$', line.strip())
            if hr_match:
                blocks.append(('hr', None))
                i += 1
                continue

            ul_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
            if ul_match:
                items = []
                indent = len(ul_match.group(1))
                while i < len(lines) and re.match(r'^(\s*)[-*+]\s+(.+)$', lines[i]):
                    m = re.match(r'^(\s*)[-*+]\s+(.+)$', lines[i])
                    items.append((len(m.group(1)), m.group(2)))
                    i += 1
                blocks.append(('unordered_list', items))
                continue

            ol_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
            if ol_match:
                items = []
                while i < len(lines) and re.match(r'^(\s*)\d+\.\s+(.+)$', lines[i]):
                    m = re.match(r'^(\s*)\d+\.\s+(.+)$', lines[i])
                    items.append((len(m.group(1)), m.group(2)))
                    i += 1
                blocks.append(('ordered_list', items))
                continue

            bq_match = re.match(r'^>\s*(.*)$', line)
            if bq_match:
                quote_lines = []
                while i < len(lines) and re.match(r'^>', lines[i]):
                    m = re.match(r'^>\s*(.*)$', lines[i])
                    quote_lines.append(m.group(1))
                    i += 1
                blocks.append(('blockquote', '\n'.join(quote_lines)))
                continue

            if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:,\-]+$', lines[i + 1]):
                headers = [h.strip() for h in line.split('|')[1:-1]]
                i += 2
                rows = []
                while i < len(lines) and '|' in lines[i]:
                    row = [c.strip() for c in lines[i].split('|')[1:-1]]
                    rows.append(row)
                    i += 1
                blocks.append(('table', (headers, rows)))
                continue

            para_lines = []
            while i < len(lines):
                if not lines[i].strip():
                    break
                if lines[i].strip().startswith('#'):
                    break
                if lines[i].strip().startswith('```'):
                    break
                if lines[i].strip().startswith('>'):
                    break
                if re.match(r'^(\s*)[-*+]\s', lines[i]):
                    break
                if re.match(r'^(\s*)\d+\.\s', lines[i]):
                    break
                if re.match(r'^[-*_]{3,}\s*$', lines[i].strip()):
                    break
                if '|' in lines[i] and i + 1 < len(lines) and re.match(r'^[\s|:,\-]+$', lines[i + 1]):
                    break
                if lines[i].strip() == '---':
                    i += 1
                    break
                para_lines.append(lines[i])
                i += 1

            if para_lines:
                text_content = '\n'.join(para_lines)
                blocks.append(('paragraph', text_content))
                continue

            i += 1

        if fence_content:
            blocks.append(('code', (fence_lang, '\n'.join(fence_content))))

        return blocks

    def _find_line_index(self, lines, start, condition):
        for i in range(start, len(lines)):
            if condition(lines[i]):
                return i
        return None

    def _extract_front_matter(self, blocks):
        metadata = {}
        for block in blocks:
            if block[0] == 'front_matter':
                try:
                    metadata = yaml.safe_load(block[1]) or {}
                except Exception:
                    pass
                break
        return metadata

    def _render_block(self, block):
        block_type = block[0]
        if block_type == 'heading':
            self._render_heading(block[1])
        elif block_type == 'paragraph':
            self._render_paragraph(block[1])
        elif block_type == 'unordered_list':
            self._render_unordered_list(block[1])
        elif block_type == 'ordered_list':
            self._render_ordered_list(block[1])
        elif block_type == 'table':
            self._render_table(block[1])
        elif block_type == 'code':
            self._render_code(block[1])
        elif block_type == 'blockquote':
            self._render_blockquote(block[1])
        elif block_type == 'hr':
            self._render_page_break()

    def _render_heading(self, heading_data):
        level, text = heading_data

        if text.lower().strip() in ('referencias', 'references', 'bibliografía', 'bibliography'):
            self.in_references = True
        else:
            self.in_references = False

        p = self.doc.add_paragraph(text, style=f'Heading {level}')
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 2.0

        for run in p.runs:
            run.font.name = 'Times New Roman'

    def _render_paragraph(self, text):
        if self.in_references:
            self._render_reference(text)
            return

        text = text.replace('\n', ' ')
        p = self.doc.add_paragraph()
        self._apply_inline_formatting(p, text)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 2.0

    def _render_reference(self, text):
        add_reference_entry(self.doc, text)

    def _apply_inline_formatting(self, paragraph, text):
        pattern = re.compile(
            r'(`+)(.+?)\1|'
            r'(!\[([^\]]*)\]\(([^)]+)\))|'
            r'(\[([^\]]*)\]\(([^)]+)\))|'
            r'(\*\*\*)(.+?)(\*\*\*)|'
            r'(\*\*)(.+?)(\*\*)|'
            r'(\*)(.+?)(\*)'
        )

        last_end = 0
        for match in pattern.finditer(text):
            start = match.start()
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                self._style_run(run)

            if match.group(1):
                run = paragraph.add_run(match.group(2))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif match.group(3):
                self._insert_image(paragraph, match.group(5), match.group(4))
            elif match.group(6):
                run = paragraph.add_run(match.group(7))
                self._style_run(run)
                run.font.underline = True
            elif match.group(9):
                run = paragraph.add_run(match.group(10))
                self._style_run(run, bold=True, italic=True)
            elif match.group(12):
                run = paragraph.add_run(match.group(13))
                self._style_run(run, bold=True)
            elif match.group(15):
                run = paragraph.add_run(match.group(16))
                self._style_run(run, italic=True)

            last_end = match.end()

        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            self._style_run(run)

    @staticmethod
    def _style_run(run, bold=False, italic=False):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    def _insert_image(self, paragraph, src, alt_text):
        try:
            if not os.path.isabs(src):
                base_dir = os.path.dirname(self.md_path)
                src = os.path.join(base_dir, src)
            if os.path.exists(src):
                run = paragraph.add_run()
                run.add_picture(src, width=Inches(4.5))
                last_run = paragraph.add_run()
                last_run.add_break()
        except Exception:
            pass

    def _render_unordered_list(self, items):
        for indent, text in items:
            level = indent // 2
            p = self.doc.add_paragraph()
            bullet_char = '●' if level == 0 else '○' if level == 1 else '■'
            r = p.add_run(f'{bullet_char} ')
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.bold = True
            self._apply_inline_formatting(p, text)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 2.0
            pf.left_indent = Inches(0.5 + 0.25 * level)

    def _render_ordered_list(self, items):
        for i, (indent, text) in enumerate(items, 1):
            level = indent // 2
            p = self.doc.add_paragraph()
            r = p.add_run(f'{i}. ')
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.bold = True
            self._apply_inline_formatting(p, text)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 2.0
            pf.left_indent = Inches(0.5 + 0.25 * level)

    def _render_table(self, table_data):
        headers, rows = table_data
        num_cols = max(len(headers), max((len(r) for r in rows), default=0))
        if num_cols == 0:
            return

        table = self.doc.add_table(rows=1 + len(rows), cols=num_cols)
        table.style = 'Table Grid'
        table.autofit = True

        for j, header in enumerate(headers):
            if j < num_cols:
                cell = table.rows[0].cells[j]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    cell = table.rows[i + 1].cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)

        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 2.0

    def _render_code(self, code_data):
        lang, code_text = code_data
        for line in code_text.split('\n'):
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
            pf.left_indent = Inches(0.5)

    def _render_blockquote(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True
        set_paragraph_format(p, first_line_indent=Inches(0.5))
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)

    def _render_page_break(self):
        self.doc.add_page_break()


def convert(md_path, output_path, include_cover=True, include_page_numbers=True):
    converter = MdToDocxConverter(md_path)
    return converter.convert(output_path, include_cover, include_page_numbers)
